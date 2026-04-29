"""
F1 Commentary Sentiment Analysis using AFINN Lexicon (Nielsen, 2011) — v3

Reads .docx transcripts from raw_content/, performs AFINN-based sentiment
scoring per driver, and outputs a structured Excel workbook with one coding
sheet per .docx file (no empty padding sheets) and two summary sheets
(F1 Live vs Sky Sports).

Each summary sheet contains:
  - An overall summary table (all races combined) at the top
  - Per-race tables arranged side by side below the charts

Dependencies: afinn, python-docx, xlsxwriter
"""

import re
import glob
from datetime import datetime
from pathlib import Path

from afinn import Afinn
from docx import Document
import xlsxwriter

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
RAW_DIR = BASE_DIR / "raw_content"
TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
OUTPUT_FILE = BASE_DIR / f"F1_Sentiment_{TIMESTAMP}.xlsx"

DRIVERS = [
    ("Oscar Piastri", "Australian"),
    ("Lando Norris", "British"),
    ("Max Verstappen", "Dutch"),
    ("George Russell", "British"),
    ("Charles Leclerc", "Monegasque"),
    ("Lewis Hamilton", "British"),
    ("Alexander Albon", "Thailand"),
    ("Kimi Antonelli", "Italian"),
    ("Isack Hadjar", "French"),
    ("Nico Hulkenberg", "German"),
    ("Lance Stroll", "Canadian"),
    ("Fernando Alonso", "Spanish"),
    ("Esteban Ocon", "French"),
    ("Pierre Gasly", "French"),
    ("Liam Lawson", "New Zealander"),
    ("Gabriel Bortoleto", "Brazilian"),
    ("Oliver Bearman", "British"),
    ("Carlos Sainz", "Spanish"),
    ("Yuki Tsunoda", "Japanese"),
    ("Franco Colapinto", "Argentinian"),
    ("Jack Doohan", "Australian"),
]

NATIONALITY_KEYWORDS = {
    "Australian": ["australian", "australia", "aussie"],  # added "aussie"
    "British": ["british", "britain", "uk", "united kingdom", "england", "english"],
    "Dutch": ["dutch", "netherlands", "holland"],
    "Monegasque": ["monegasque", "monaco"],
    "Thailand": ["thai", "thailand"],
    "Italian": ["italian", "italy"],
    "French": ["french", "france"],
    "German": ["german", "germany"],
    "Canadian": ["canadian", "canada"],
    "Spanish": ["spanish", "spain"],
    "New Zealander": ["new zealand", "kiwi"],
    "Brazilian": ["brazilian", "brazil"],
    "Japanese": ["japanese", "japan"],
    "Argentinian": ["argentinian", "argentina", "argentine"],
}

# Phrases that should not contribute to AFINN scoring or highlighting.
# Replaced with equal-length spaces before any AFINN operation to preserve
# character offsets for rich-string highlighting.
# CHANGE 1 (lines 75-86 in v2):
#   - Added "british grand prix" and "british gp" BEFORE "grand prix" so that
#     the event name is masked before the word "british" is checked for
#     nationality. More specific phrases must always come first.
#   - Added "dirty air" (racing term, not a sentiment word).
#   - Added "stop" (neutral race action).
NEUTRAL_PHRASES = [
    "british grand prix",   # must precede "grand prix" — event name, not nationality
    "british gp",           # same reason
    "grand prix",
    "safety car",
    "dirty air",            # NEW — racing term, not a sentiment expression
    "hard tire",
    "hard tyre",
    "hard tires",
    "hard tyres",
    "slick",
    "yeah",
    "no",
    "stop",                 # NEW — neutral race reporting action
    "crash",
]

# Extra phrases that trigger a mention for a specific driver.
# List longer/more-specific aliases before shorter ones (regex tries left-to-right).
# CHANGE 2 (lines 90-95 in v2):
#   - Added "Ollie" as a common alias for Oliver Bearman so any mention of
#     "Ollie" in transcripts is attributed to him automatically.
DRIVER_ALIASES = {
    "Max Verstappen": [
        "reigning world champion",  # more specific first
        "reigning champion",
    ],
    "Oliver Bearman": [     # NEW
        "Ollie",            # common nickname used in commentary
    ],
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

afinn = Afinn()

CUSTOM_SCORES = {
    'attack': 3, 'attacks': 3, 'attacking': 3, 'attacked': 3, 'attacker': 3,
    'mesmeric': 5,
    'ahead': 2,
    'faster': 2,
    'slower': -2,
}
afinn._dict.update(CUSTOM_SCORES)
afinn._pattern = re.compile(
    r'\b(?:' + '|'.join(re.escape(w) for w in sorted(afinn._dict, key=len, reverse=True)) + r')\b'
)


def extract_race_name(doc: Document, filename: str) -> str:
    for p in doc.paragraphs[:3]:
        text = p.text.strip()
        if text and "GP" in text:
            return text
    return Path(filename).stem


def build_driver_patterns(drivers):
    patterns = {}
    for full_name, nationality in drivers:
        parts = full_name.split()
        first, last = parts[0], parts[-1]
        name_tokens = [re.escape(first), re.escape(last), re.escape(full_name)]

        alias_tokens = [
            re.escape(alias)
            for alias in DRIVER_ALIASES.get(full_name, [])
        ]

        all_tokens = name_tokens + alias_tokens
        pattern = re.compile(
            r"\b(?:" + "|".join(all_tokens) + r")\b", re.IGNORECASE
        )
        patterns[full_name] = pattern
    return patterns


def split_into_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in parts if s.strip()]


def find_mentions(paragraphs: list[str], pattern: re.Pattern) -> list[str]:
    mentions = []
    for para in paragraphs:
        for sent in split_into_sentences(para):
            if pattern.search(sent):
                mentions.append(sent)
    return mentions


def check_nationality_mention(mentions: list[str], nationality: str) -> str:
    keywords = NATIONALITY_KEYWORDS.get(nationality, [nationality.lower()])
    for mention in mentions:
        text_lower = mention.lower()
        if any(kw in text_lower for kw in keywords):
            return "Yes"
    return "No"


def count_nationality_mentions(mentions: list[str], nationality: str) -> int:
    keywords = NATIONALITY_KEYWORDS.get(nationality, [nationality.lower()])
    count = 0
    for mention in mentions:
        text_lower = mention.lower()
        for kw in keywords:
            count += text_lower.count(kw)
    return count


def _is_in_parentheses(text: str, start: int, end: int) -> bool:
    """Return True if the match at [start:end] is directly wrapped in parentheses.

    Checks that the nearest non-space character before `start` is '('
    and the nearest non-space character after `end` is ')'.
    """
    before = text[:start].rstrip()
    after  = text[end:].lstrip()
    return bool(before) and before[-1] == "(" and bool(after) and after[0] == ")"


def count_name_mentions(paragraphs: list[str], pattern: re.Pattern, full_name: str) -> int:
    """Count distinct driver name mentions across all paragraphs.

    Two fixes vs v3:
      1. Deduplication: "Lando Norris" counts as 1, not 2 (first + last).
         Full name matches are recorded first; any first/last match whose span
         overlaps a full name match is skipped.
      2. Parenthesis exclusion: mentions like (Norris) contribute to AFINN
         scoring but are NOT counted as explicit name mentions for frequency.
    """
    full_name_pattern = re.compile(re.escape(full_name), re.IGNORECASE)

    total = 0
    for para in paragraphs:
        # Collect spans already claimed by full-name matches
        full_spans = [(m.start(), m.end()) for m in full_name_pattern.finditer(para)]

        # Each full name match = 1 mention, unless it is in parentheses
        for fs, fe in full_spans:
            if not _is_in_parentheses(para, fs, fe):
                total += 1

        # Check all other matches (first name, last name, aliases)
        for m in pattern.finditer(para):
            start, end = m.start(), m.end()

            # Skip if this span is already covered by a full-name match
            if any(fs <= start and end <= fe for fs, fe in full_spans):
                continue

            # Skip if wrapped in parentheses
            if _is_in_parentheses(para, start, end):
                continue

            total += 1

    return total


def classify_sentiment(score: float) -> str:
    if score > 0:
        return "Positive"
    elif score < 0:
        return "Negative"
    return "Neutral"


def mask_neutral_phrases(text: str) -> str:
    result = text
    for phrase in NEUTRAL_PHRASES:
        result = re.sub(re.escape(phrase), " " * len(phrase), result, flags=re.IGNORECASE)
    return result


def score_attributed_for_driver(sentence: str, target_driver: str, all_patterns: dict) -> float:
    """AFINN score for a sentence, attributing only words closest to target_driver."""
    driver_positions = {}
    for d_name, pat in all_patterns.items():
        for m in pat.finditer(sentence):
            driver_positions.setdefault(d_name, []).append((m.start() + m.end()) / 2.0)

    other_drivers = [d for d in driver_positions if d != target_driver]
    if not other_drivers:
        return afinn.score(mask_neutral_phrases(sentence))

    all_mentions = [(pos, d) for d, positions in driver_positions.items() for pos in positions]
    masked = mask_neutral_phrases(re.sub(r"\s+", " ", sentence))
    total = 0.0
    for match in afinn._pattern.finditer(masked.lower()):
        word_score = afinn._dict.get(match.group(), 0)
        if word_score == 0:
            continue
        word_mid = (match.start() + match.end()) / 2.0
        nearest = min(all_mentions, key=lambda x: abs(x[0] - word_mid))[1]
        if nearest == target_driver:
            total += word_score
    return total


def get_afinn_word_counts(text: str) -> dict:
    """Count positive, negative, and neutral AFINN words in text.

    Returns dict with keys: pos_count, neg_count, neu_count,
    pos_words, neg_words, neu_words (lists of (word, score) tuples).
    """
    normalized = re.sub(r"\s+", " ", text)
    masked = mask_neutral_phrases(normalized)
    lowered = masked.lower()
    matches = list(afinn._pattern.finditer(lowered))

    result = {
        "pos_count": 0, "neg_count": 0, "neu_count": 0,
        "pos_words": [], "neg_words": [], "neu_words": [],
    }
    for match in matches:
        word = match.group()
        score = afinn._dict.get(word, 0)
        if score > 0:
            result["pos_count"] += 1
            result["pos_words"].append((word, score))
        elif score < 0:
            result["neg_count"] += 1
            result["neg_words"].append((word, score))
        else:
            result["neu_count"] += 1
            result["neu_words"].append((word, score))
    return result


def get_afinn_word_counts_attributed(sentences: list[str], target_driver: str, all_patterns: dict) -> dict:
    """Like get_afinn_word_counts but applies proximity logic to attribute words to target_driver."""
    result = {"pos_count": 0, "neg_count": 0, "neu_count": 0,
              "pos_words": [], "neg_words": [], "neu_words": []}
    for sentence in sentences:
        driver_positions = {}
        for d_name, pat in all_patterns.items():
            for m in pat.finditer(sentence):
                driver_positions.setdefault(d_name, []).append((m.start() + m.end()) / 2.0)
        other_drivers = [d for d in driver_positions if d != target_driver]
        all_mentions = [(pos, d) for d, positions in driver_positions.items() for pos in positions]
        masked = mask_neutral_phrases(re.sub(r"\s+", " ", sentence))
        for match in afinn._pattern.finditer(masked.lower()):
            word = match.group()
            score = afinn._dict.get(word, 0)
            if other_drivers and all_mentions:
                word_mid = (match.start() + match.end()) / 2.0
                nearest = min(all_mentions, key=lambda x: abs(x[0] - word_mid))[1]
                if nearest != target_driver:
                    continue
            if score > 0:
                result["pos_count"] += 1; result["pos_words"].append((word, score))
            elif score < 0:
                result["neg_count"] += 1; result["neg_words"].append((word, score))
            else:
                result["neu_count"] += 1; result["neu_words"].append((word, score))
    return result


def build_rich_string_args(text: str, green_fmt, red_fmt, black_fmt) -> list:
    """Build argument list for xlsxwriter's write_rich_string().

    Returns alternating [format, string, ...] where:
      - Positive AFINN words (score > 0) use green_fmt
      - Negative AFINN words (score < 0) use red_fmt
      - Normal text uses black_fmt
    Uses AFINN's own internal pattern for exact match fidelity.
    """
    normalized = re.sub(r"\s+", " ", text)
    masked = mask_neutral_phrases(normalized)
    lowered = masked.lower()
    matches = list(afinn._pattern.finditer(lowered))

    if not matches:
        return []  # Signal: no AFINN words, use plain write()

    parts = []
    last_end = 0

    for match in matches:
        start, end = match.start(), match.end()
        word = match.group()
        score = afinn._dict.get(word, 0)

        if start > last_end:
            parts.append(black_fmt)
            parts.append(normalized[last_end:start])

        if score > 0:
            parts.append(green_fmt)
        elif score < 0:
            parts.append(red_fmt)
        else:
            parts.append(black_fmt)
        parts.append(normalized[start:end])
        last_end = end

    if last_end < len(normalized):
        parts.append(black_fmt)
        parts.append(normalized[last_end:])

    return parts


# CHANGE 3 (replaces write_summary_sheet lines 326-483 in v2):
# Restructured to support the new layout:
#   - Overall summary table at the top (rows 1..N) — unchanged from v2
#   - 3 charts below the overall table — unchanged from v2
#   - Per-race tables side by side below the charts, one column-block per race
#     Each block has a race title header row, then column headers, then driver rows.
# Also accepts a new `race_names` parameter (ordered list of race names for
# this broadcast group) so the per-race tables are labelled correctly.
def write_summary_sheet(wb, sheet_title: str, summary_data: dict, wb_formats: dict,
                        race_names: list[str]):
    """Write a summary worksheet with overall table, charts, and per-race tables."""
    ws_sum = wb.add_worksheet(sheet_title)

    header_fmt        = wb_formats["header_fmt"]
    cell_fmt          = wb_formats["cell_fmt"]
    pos_score_fmt     = wb_formats["pos_score_fmt"]
    neg_score_fmt     = wb_formats["neg_score_fmt"]
    neutral_score_fmt = wb_formats["neutral_score_fmt"]
    pos_word_fmt      = wb_formats["pos_word_fmt"]
    neg_word_fmt      = wb_formats["neg_word_fmt"]
    race_title_fmt    = wb_formats["race_title_fmt"]

    # ------------------------------------------------------------------
    # SECTION 1: Overall summary table (top, row 0 = headers, rows 1..N = drivers)
    # Identical to v2 — no logic changes here.
    # ------------------------------------------------------------------
    sum_headers = [
        "Driver", "Avg AFINN Score", "% Positive",
        "% Negative", "% Neutral",
        "Positive Words", "Negative Words", "Neutral Words", "Total AFINN Words",
        "Nationality Mentioned Count",
        "Name Mention Count",
    ]
    sum_widths = [22, 18, 14, 14, 14, 16, 16, 16, 18, 26, 20]
    for c, (h, w) in enumerate(zip(sum_headers, sum_widths)):
        ws_sum.write(0, c, h, header_fmt)
        ws_sum.set_column(c, c, w)

    print(f"\n{'=' * 60}")
    print(f"SUMMARY — {sheet_title}")
    print("=" * 60)

    for drv_idx, (driver_name, nationality) in enumerate(DRIVERS):
        row = drv_idx + 1
        data = summary_data[driver_name]
        scores = data["scores"]

        ws_sum.write(row, 0, driver_name, cell_fmt)

        if scores:
            avg_score = sum(scores) / len(scores)
            total = len(scores)
            pct_pos = round(sum(1 for s in scores if s > 0) / total * 100, 1)
            pct_neg = round(sum(1 for s in scores if s < 0) / total * 100, 1)
            pct_neu = round(sum(1 for s in scores if s == 0) / total * 100, 1)
        else:
            avg_score = 0
            pct_pos = pct_neg = pct_neu = 0

        pw  = data["pos_words_total"]
        nw  = data["neg_words_total"]
        nuw = data["neu_words_total"]
        tw  = pw + nw + nuw

        score_fmt = (
            pos_score_fmt if avg_score > 0
            else neg_score_fmt if avg_score < 0
            else neutral_score_fmt
        )
        ws_sum.write(row, 1, round(avg_score, 2), score_fmt)
        ws_sum.write(row, 2, pct_pos, cell_fmt)
        ws_sum.write(row, 3, pct_neg, cell_fmt)
        ws_sum.write(row, 4, pct_neu, cell_fmt)
        ws_sum.write(row, 5, pw,  pos_word_fmt)
        ws_sum.write(row, 6, nw,  neg_word_fmt)
        ws_sum.write(row, 7, nuw, cell_fmt)
        ws_sum.write(row, 8, tw,  cell_fmt)
        ws_sum.write(row, 9, data["nat_count"],     cell_fmt)
        ws_sum.write(row, 10, data["mention_count"], cell_fmt)

        print(
            f"  {driver_name:22s} | Avg: {avg_score:7.2f} | "
            f"+{pct_pos:5.1f}% -{pct_neg:5.1f}% ={pct_neu:5.1f}% | "
            f"Words [+{pw} -{nw} ={nuw} T:{tw}] | "
            f"Nat: {data['nat_count']} | Names: {data['mention_count']}"
        )

    num_drivers = len(DRIVERS)

    # ------------------------------------------------------------------
    # SECTION 2: Charts (identical to v2, positioned below overall table)
    # ------------------------------------------------------------------

    # Chart 1: Average AFINN score per driver
    chart1 = wb.add_chart({"type": "column"})
    chart1.set_title({"name": f"Average AFINN Sentiment Score per Driver — {sheet_title}"})
    chart1.set_y_axis({"name": "AFINN Score"})
    chart1.set_x_axis({"name": "Driver"})
    chart1.set_size({"width": 960, "height": 480})
    chart1.add_series({
        "name": "Avg AFINN Score",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 1, num_drivers, 1],
        "points": [
            {
                "fill": {
                    "color": (
                        "#27AE60" if (
                            sum(summary_data[d]["scores"]) / len(summary_data[d]["scores"])
                            if summary_data[d]["scores"] else 0
                        ) > 0
                        else "#E74C3C" if (
                            sum(summary_data[d]["scores"]) / len(summary_data[d]["scores"])
                            if summary_data[d]["scores"] else 0
                        ) < 0
                        else "#F39C12"
                    )
                }
            }
            for d, _ in DRIVERS
        ],
    })
    chart1.set_legend({"none": True})
    ws_sum.insert_chart("A" + str(num_drivers + 3), chart1)

    # Chart 2: Stacked sentiment distribution
    chart2 = wb.add_chart({"type": "column", "subtype": "stacked"})
    chart2.set_title({"name": f"Sentiment Distribution per Driver — {sheet_title}"})
    chart2.set_y_axis({"name": "Percentage (%)"})
    chart2.set_x_axis({"name": "Driver"})
    chart2.set_size({"width": 960, "height": 480})
    chart2.add_series({
        "name": "% Positive",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 2, num_drivers, 2],
        "fill": {"color": "#27AE60"},
    })
    chart2.add_series({
        "name": "% Negative",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 3, num_drivers, 3],
        "fill": {"color": "#E74C3C"},
    })
    chart2.add_series({
        "name": "% Neutral",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 4, num_drivers, 4],
        "fill": {"color": "#F39C12"},
    })
    ws_sum.insert_chart("A" + str(num_drivers + 20), chart2)

    # Chart 3: Word count comparison per driver
    chart3 = wb.add_chart({"type": "column", "subtype": "stacked"})
    chart3.set_title({"name": f"AFINN Word Count per Driver — {sheet_title}"})
    chart3.set_y_axis({"name": "Word Count"})
    chart3.set_x_axis({"name": "Driver"})
    chart3.set_size({"width": 960, "height": 480})
    chart3.add_series({
        "name": "Positive Words",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 5, num_drivers, 5],
        "fill": {"color": "#27AE60"},
    })
    chart3.add_series({
        "name": "Negative Words",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 6, num_drivers, 6],
        "fill": {"color": "#E74C3C"},
    })
    chart3.add_series({
        "name": "Neutral Words",
        "categories": [sheet_title, 1, 0, num_drivers, 0],
        "values":     [sheet_title, 1, 7, num_drivers, 7],
        "fill": {"color": "#F39C12"},
    })
    ws_sum.insert_chart("A" + str(num_drivers + 37), chart3)

    # ------------------------------------------------------------------
    # SECTION 3 (NEW): Per-race tables, side by side below the charts.
    #
    # Layout:
    #   - Each race occupies a block of columns: 8 data columns + 1 spacer
    #   - All blocks share the same row range so they sit side by side
    #   - Row layout within each block:
    #       per_race_start_row + 0  = race title (merged across 8 cols)
    #       per_race_start_row + 1  = column headers
    #       per_race_start_row + 2  .. +2+N = one row per driver
    #
    # per_race_start_row is placed after the last chart (chart3 ends at
    # num_drivers+37, each chart is ~30 rows tall at default scale, so
    # we use num_drivers + 37 + 32 as a safe offset).
    # ------------------------------------------------------------------

    per_race_headers = [
        "Driver", "Avg Score", "% Pos", "% Neg", "% Neu",
        "Pos Words", "Neg Words", "Nat Count",
    ]
    # Number of columns per race block (data cols + 1 spacer col)
    block_width = len(per_race_headers) + 1

    # Start row for the per-race section (below all three charts)
    per_race_start_row = num_drivers + 37 + 32

    for race_idx, race_name in enumerate(race_names):
        # Starting column for this race block
        col_start = race_idx * block_width

        # Row 0 of block: race title spanning all data columns
        ws_sum.merge_range(
            per_race_start_row, col_start,
            per_race_start_row, col_start + len(per_race_headers) - 1,
            race_name, race_title_fmt,
        )

        # Row 1 of block: column headers
        for c, h in enumerate(per_race_headers):
            ws_sum.write(per_race_start_row + 1, col_start + c, h, header_fmt)

        # Rows 2..N+2: one row per driver
        for drv_idx, (driver_name, _) in enumerate(DRIVERS):
            data_row = per_race_start_row + 2 + drv_idx
            race_entry = summary_data[driver_name]["per_race"].get(race_name)

            ws_sum.write(data_row, col_start, driver_name, cell_fmt)

            if race_entry and race_entry["scores"]:
                r_scores = race_entry["scores"]
                r_avg    = sum(r_scores) / len(r_scores)
                r_total  = len(r_scores)
                r_ppos   = round(sum(1 for s in r_scores if s > 0) / r_total * 100, 1)
                r_pneg   = round(sum(1 for s in r_scores if s < 0) / r_total * 100, 1)
                r_pneu   = round(sum(1 for s in r_scores if s == 0) / r_total * 100, 1)
                r_pw     = race_entry["pos_words_total"]
                r_nw     = race_entry["neg_words_total"]
                r_nat    = race_entry["nat_count"]

                r_score_fmt = (
                    pos_score_fmt if r_avg > 0
                    else neg_score_fmt if r_avg < 0
                    else neutral_score_fmt
                )
                ws_sum.write(data_row, col_start + 1, round(r_avg, 2), r_score_fmt)
                ws_sum.write(data_row, col_start + 2, r_ppos, cell_fmt)
                ws_sum.write(data_row, col_start + 3, r_pneg, cell_fmt)
                ws_sum.write(data_row, col_start + 4, r_pneu, cell_fmt)
                ws_sum.write(data_row, col_start + 5, r_pw,   pos_word_fmt)
                ws_sum.write(data_row, col_start + 6, r_nw,   neg_word_fmt)
                ws_sum.write(data_row, col_start + 7, r_nat,  cell_fmt)
            else:
                # Driver had no mentions in this race — write blanks
                for c in range(1, len(per_race_headers)):
                    ws_sum.write(data_row, col_start + c, "", cell_fmt)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

# CHANGE 4 (lines 490-495 in v2):
# Added "per_race" dict to each driver entry so race-by-race data can be
# stored alongside the cumulative totals.
def _empty_driver_entry():
    return {
        "scores": [], "nat_count": 0,
        "pos_words_total": 0, "neg_words_total": 0, "neu_words_total": 0,
        "mention_count": 0,
        "per_race": {},          # NEW — keyed by race_name, same structure minus per_race
    }


def _empty_race_entry():
    """Per-race accumulator stored inside driver["per_race"][race_name]."""
    return {
        "scores": [],
        "pos_words_total": 0, "neg_words_total": 0, "neu_words_total": 0,
        "nat_count": 0, "mention_count": 0,
    }


def process_documents():
    docx_files = sorted(glob.glob(str(RAW_DIR / "*.docx")))
    if not docx_files:
        print(f"ERROR: No .docx files found in {RAW_DIR}")
        return

    print(f"Found {len(docx_files)} document(s) in {RAW_DIR}\n")

    driver_patterns = build_driver_patterns(DRIVERS)
    wb = xlsxwriter.Workbook(str(OUTPUT_FILE))

    # --- Formats ---
    header_fmt = wb.add_format({
        "bold": True, "font_size": 11, "font_color": "white",
        "bg_color": "#4472C4", "border": 1,
        "align": "center", "valign": "vcenter",
    })
    cell_fmt = wb.add_format({"border": 1, "valign": "top"})
    cell_center = wb.add_format({"border": 1, "align": "center", "valign": "top"})
    cell_wrap = wb.add_format({"border": 1, "text_wrap": True, "valign": "top"})
    pos_fmt = wb.add_format({
        "border": 1, "align": "center", "bg_color": "#C6EFCE",
    })
    neg_fmt = wb.add_format({
        "border": 1, "align": "center", "bg_color": "#FFC7CE",
    })
    neu_fmt = wb.add_format({
        "border": 1, "align": "center", "bg_color": "#FFEB9C",
    })
    # Rich text formats (no cell-level properties, just font)
    red_font = wb.add_format({"font_color": "#E74C3C", "bold": True})
    green_font = wb.add_format({"font_color": "#27AE60", "bold": True})
    black_font = wb.add_format({"font_color": "black"})

    # Summary score/word formats
    pos_score_fmt     = wb.add_format({"border": 1, "bg_color": "#C6EFCE", "num_format": "0.00"})
    neg_score_fmt     = wb.add_format({"border": 1, "bg_color": "#FFC7CE", "num_format": "0.00"})
    neutral_score_fmt = wb.add_format({"border": 1, "num_format": "0.00"})
    pos_word_fmt      = wb.add_format({"border": 1, "font_color": "#27AE60", "bold": True})
    neg_word_fmt      = wb.add_format({"border": 1, "font_color": "#E74C3C", "bold": True})

    # NEW: Race title format for per-race table headers
    race_title_fmt = wb.add_format({
        "bold": True, "font_size": 12, "font_color": "white",
        "bg_color": "#2E4057", "border": 1,
        "align": "center", "valign": "vcenter",
    })

    wb_formats = {
        "header_fmt": header_fmt, "cell_fmt": cell_fmt,
        "pos_score_fmt": pos_score_fmt, "neg_score_fmt": neg_score_fmt,
        "neutral_score_fmt": neutral_score_fmt,
        "pos_word_fmt": pos_word_fmt, "neg_word_fmt": neg_word_fmt,
        "race_title_fmt": race_title_fmt,   # NEW — passed into write_summary_sheet
    }

    # CHANGE 5 (lines 547-548 in v2):
    # Accumulators now use the updated _empty_driver_entry() which includes
    # the "per_race" dict. Also added two ordered lists to track which race
    # names belong to each broadcast group — needed to label per-race tables.
    summary_live = {name: _empty_driver_entry() for name, _ in DRIVERS}
    summary_sky  = {name: _empty_driver_entry() for name, _ in DRIVERS}
    race_names_live = []   # NEW — ordered list of race names for F1 Live
    race_names_sky  = []   # NEW — ordered list of race names for Sky Sports

    headers = [
        "Drivers", "Nationality",
        "Language used (positive, neutral, negative)",
        "Quotes/Transcript", "Mention of nationality (yes/no)",
        "Nationality Mention Count",
        "Name Mention Count",
    ]
    col_widths = [22, 16, 18, 80, 18, 22, 20]

    for race_idx, docx_path in enumerate(docx_files):
        doc = Document(docx_path)
        race_name = extract_race_name(doc, docx_path)

        # Sheet name derived from the .docx filename (Excel sheet names max 31 chars)
        sheet_name = re.sub(r'[\[\]:*?/\\]', '_', Path(docx_path).stem)[:31]

        # Determine broadcast group for summary accumulation
        is_sky = 'sky' in Path(docx_path).stem.lower()
        summary_data = summary_sky if is_sky else summary_live

        # CHANGE 5b: Track race names per broadcast group for the per-race tables
        if is_sky:
            if race_name not in race_names_sky:
                race_names_sky.append(race_name)
        else:
            if race_name not in race_names_live:
                race_names_live.append(race_name)

        print(f"{'='*60}")
        print(f"Processing: {race_name}")
        print(f"  File: {Path(docx_path).name}")
        print(f"  Sheet: {sheet_name}  ({'Sky Sports' if is_sky else 'F1 Live'})")
        print(f"{'='*60}")

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        ws = wb.add_worksheet(sheet_name)

        for c, (h, w) in enumerate(zip(headers, col_widths)):
            ws.write(0, c, h, header_fmt)
            ws.set_column(c, c, w)

        drivers_found = []

        for drv_idx, (driver_name, nationality) in enumerate(DRIVERS):
            row = drv_idx + 1
            ws.write(row, 0, driver_name, cell_fmt)
            ws.write(row, 1, nationality, cell_fmt)

            pattern = driver_patterns[driver_name]
            mentions = find_mentions(paragraphs, pattern)

            if mentions:
                # Attributed scoring: words closest to target driver get counted
                score = sum(
                    score_attributed_for_driver(s, driver_name, driver_patterns)
                    for s in mentions
                )
                sentiment = classify_sentiment(score)

                # Column C — sentiment
                sent_format = (
                    pos_fmt if sentiment == "Positive"
                    else neg_fmt if sentiment == "Negative"
                    else neu_fmt
                )
                ws.write(row, 2, sentiment, sent_format)

                # Column D — quotes with AFINN words colored (green=pos, red=neg)
                combined_text = " ".join(mentions)
                quotes = " | ".join(mentions)
                rich_args = build_rich_string_args(quotes, green_font, red_font, black_font)
                if rich_args:
                    ws.write_rich_string(row, 3, *rich_args, cell_wrap)
                else:
                    ws.write(row, 3, quotes, cell_wrap)

                # Column E — nationality mention
                nat_check = check_nationality_mention(mentions, nationality)
                ws.write(row, 4, nat_check, cell_center)

                # Column F — exact nationality mention count
                nat_count_race = count_nationality_mentions(mentions, nationality)
                ws.write(row, 5, nat_count_race, cell_center)

                # Column G — driver name mention count (across all paragraphs)
                name_count = count_name_mentions(paragraphs, pattern, driver_name)
                ws.write(row, 6, name_count, cell_center)

                # Word counts for summary (attributed)
                word_counts = get_afinn_word_counts_attributed(mentions, driver_name, driver_patterns)

                # Cumulative totals (same as v2)
                summary_data[driver_name]["scores"].append(score)
                summary_data[driver_name]["pos_words_total"] += word_counts["pos_count"]
                summary_data[driver_name]["neg_words_total"] += word_counts["neg_count"]
                summary_data[driver_name]["neu_words_total"] += word_counts["neu_count"]
                summary_data[driver_name]["nat_count"]       += nat_count_race
                summary_data[driver_name]["mention_count"]   += name_count

                # CHANGE 6 (lines 630-638 in v2):
                # Also write the same data into per_race so each race can be
                # shown separately in the per-race tables on the summary sheet.
                if race_name not in summary_data[driver_name]["per_race"]:
                    summary_data[driver_name]["per_race"][race_name] = _empty_race_entry()
                pr = summary_data[driver_name]["per_race"][race_name]
                pr["scores"].append(score)
                pr["pos_words_total"] += word_counts["pos_count"]
                pr["neg_words_total"] += word_counts["neg_count"]
                pr["nat_count"]       += nat_count_race
                pr["mention_count"]   += name_count

                drivers_found.append(
                    f"    {driver_name}: score={score:.1f} ({sentiment}), "
                    f"{len(mentions)} mention(s), nationality={nat_check} ({nat_count_race}x), "
                    f"name mentions={name_count}"
                )
            else:
                ws.write(row, 2, "", cell_fmt)
                ws.write(row, 3, "", cell_fmt)
                ws.write(row, 4, "", cell_fmt)
                ws.write(row, 5, "", cell_fmt)
                ws.write(row, 6, 0, cell_center)

        print(f"  Drivers found: {len(drivers_found)}")
        for line in drivers_found:
            print(line)
        print()

    # CHANGE 7 (lines 657-659 in v2):
    # Pass race_names lists into write_summary_sheet so it can render the
    # per-race tables with the correct race labels.
    write_summary_sheet(wb, "Summary_F1_Live",      summary_live, wb_formats, race_names_live)
    write_summary_sheet(wb, "Summary_F1_Sky_Sports", summary_sky,  wb_formats, race_names_sky)

    wb.close()
    print(f"\nOutput saved to: {OUTPUT_FILE}")
    print("\n" + "=" * 60)
    print("WHAT THIS SCRIPT GENERATES")
    print("=" * 60)
    print(
        "  An Excel workbook with:\n"
        "  - One coding sheet per .docx file (named after the file)\n"
        "  - No empty padding sheets — only real data\n"
        "  - Columns A-G: Driver, Nationality, Sentiment, Quotes (color-coded),\n"
        "                 Nationality mention (yes/no), Nationality count, Name mention count\n"
        "  - Summary_F1_Live sheet:\n"
        "      * Overall summary table at the top\n"
        "      * 3 charts (avg score, sentiment %, word counts)\n"
        "      * Per-race tables side by side below the charts\n"
        "  - Summary_F1_Sky_Sports sheet: same structure as above\n"
        "  - Neutral phrases now include: british grand prix, british gp,\n"
        "    dirty air, stop (in addition to existing list)\n"
        "  - 'Ollie' is recognised as an alias for Oliver Bearman\n"
    )
    print("REQUIREMENTS")
    print("  pip install afinn python-docx xlsxwriter")
    print("\nTIP: Use collections.Counter on name mentions for frequency analysis:")
    print("  from collections import Counter")
    print("  # Counter({name: summary['mention_count'] for name, summary in data.items()})")
    print("Done!")


if __name__ == "__main__":
    process_documents()
